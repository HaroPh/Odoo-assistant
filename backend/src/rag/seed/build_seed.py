import os
from docx import Document
from openpyxl import Workbook

HERE = os.path.dirname(__file__)


def _docx(name, title, sections):
    d = Document()
    d.add_heading(title, level=1)
    for head, body in sections:
        d.add_heading(head, level=2)
        d.add_paragraph(body)
    d.save(os.path.join(HERE, name))


def build():
    _docx("policy.docx", "Chính sách hoàn hàng", [
        ("Mục 1 — Điều kiện hoàn hàng",
         "Khách hàng có thể hoàn hàng trong vòng 30 ngày kể từ ngày mua, "
         "sản phẩm còn nguyên vẹn và có hóa đơn. Sản phẩm phải trong tình trạng "
         "chưa qua sử dụng và còn nguyên tem mác."),
        ("Mục 2 — Ngoại lệ không được hoàn trả",
         "Hàng giảm giá không được hoàn trả. Hàng điện tử áp dụng bảo hành riêng. "
         "Thực phẩm và đồ uống không được hoàn trả sau khi mở nắp."),
        ("Mục 3 — Quy trình hoàn hàng",
         "Khách hàng liên hệ bộ phận chăm sóc khách hàng qua email hoặc hotline. "
         "Nhân viên xác nhận điều kiện hoàn hàng và cấp phiếu hoàn hàng (RMA). "
         "Khách gửi hàng về kho trong vòng 7 ngày sau khi được cấp phiếu."),
        ("Mục 4 — Hoàn tiền",
         "Sau khi nhận và kiểm tra hàng hoàn, bộ phận kế toán sẽ hoàn tiền trong "
         "vòng 5 đến 10 ngày làm việc. Hoàn tiền qua hình thức thanh toán ban đầu "
         "hoặc chuyển khoản ngân hàng theo yêu cầu khách hàng."),
        ("Mục 5 — Đổi hàng",
         "Khách hàng có thể đổi hàng thay vì hoàn tiền trong vòng 30 ngày. "
         "Sản phẩm đổi phải có giá trị tương đương hoặc cao hơn. "
         "Chênh lệch giá trị sẽ được tính thêm hoặc hoàn lại tùy trường hợp."),
    ])

    _docx("sop.docx", "Quy trình nhập kho", [
        ("Bước 1 — Kiểm đếm hàng hóa",
         "Nhân viên kho kiểm đếm số lượng thực nhận so với phiếu giao. "
         "Ghi chép đầy đủ số lượng từng mặt hàng vào biên bản nhập kho. "
         "Chụp ảnh lô hàng trước khi ký nhận để có bằng chứng."),
        ("Bước 2 — Đối chiếu với đơn mua",
         "Đối chiếu với đơn mua hàng (PO) trong hệ thống Odoo; chênh lệch phải lập biên bản. "
         "Nếu số lượng thực nhận ít hơn đơn, lập biên bản thiếu hàng và thông báo nhà cung cấp. "
         "Nếu nhận thừa, giữ lại chờ xác nhận từ phòng mua hàng."),
        ("Bước 3 — Kiểm tra chất lượng",
         "Bộ phận QC kiểm tra mẫu ngẫu nhiên tối thiểu 10% lô hàng. "
         "Hàng không đạt yêu cầu chất lượng được để riêng và báo cáo trong 24 giờ. "
         "Sản phẩm lỗi được gửi trả nhà cung cấp kèm phiếu trả hàng."),
        ("Bước 4 — Cập nhật hệ thống",
         "Sau khi kiểm tra xong, nhân viên cập nhật số lượng tồn kho vào Odoo. "
         "In nhãn mã vạch cho từng mặt hàng và dán vào vị trí lưu kho tương ứng. "
         "Lưu biên bản nhập kho vào hồ sơ điện tử theo mã lô hàng."),
        ("Bước 5 — Bàn giao",
         "Nhân viên kho ký bàn giao với đơn vị vận chuyển. "
         "Bản sao biên bản nhập kho được gửi cho phòng kế toán trong ngày. "
         "Hàng được sắp xếp vào đúng vị trí trong kho theo sơ đồ kho hiện hành."),
    ])

    _docx("sla.docx", "Thỏa thuận mức dịch vụ nhà cung cấp", [
        ("Điều 1 — Phạm vi áp dụng",
         "Thỏa thuận này áp dụng cho tất cả nhà cung cấp cấp 1 của công ty. "
         "Các điều khoản có hiệu lực từ ngày ký hợp đồng và được xem xét hàng năm. "
         "Nhà cung cấp phải đạt tất cả các tiêu chí để duy trì trạng thái đối tác ưu tiên."),
        ("Điều 2 — Chất lượng sản phẩm",
         "Tỷ lệ sản phẩm lỗi không vượt quá 2% mỗi lô hàng. "
         "Nhà cung cấp phải có chứng nhận ISO 9001 hoặc tương đương. "
         "Kiểm định chất lượng bởi bên thứ ba được thực hiện mỗi 6 tháng."),
        ("Điều 3 — Thời gian giao hàng",
         "Nhà cung cấp giao hàng trong 7 ngày làm việc kể từ ngày xác nhận đơn hàng. "
         "Đơn hàng khẩn cấp được xử lý trong 3 ngày làm việc với phụ phí 15%. "
         "Nhà cung cấp phải thông báo trước 48 giờ nếu không thể đảm bảo thời hạn giao."),
        ("Điều 4 — Đóng gói và vận chuyển",
         "Hàng hóa phải được đóng gói theo tiêu chuẩn của công ty. "
         "Nhà cung cấp chịu trách nhiệm bảo hiểm hàng hóa trong quá trình vận chuyển. "
         "Mỗi kiện hàng phải có nhãn mã vạch theo định dạng GS1-128."),
        ("Điều 5 — Phạt chậm trễ giao hàng",
         "Chậm mỗi ngày phạt 0,5% giá trị đơn hàng, tối đa 10% tổng giá trị. "
         "Phạt được khấu trừ trực tiếp vào hóa đơn thanh toán tiếp theo. "
         "Trường hợp bất khả kháng phải có văn bản xác nhận từ cơ quan có thẩm quyền."),
        ("Điều 6 — Thanh toán",
         "Điều khoản thanh toán net 30 ngày sau khi nhận hóa đơn hợp lệ. "
         "Chiết khấu 2% áp dụng nếu thanh toán trong vòng 10 ngày. "
         "Hóa đơn phải ghi đúng mã đơn hàng và mã nhà cung cấp."),
    ])

    wb = Workbook()
    ws = wb.active
    ws.title = "Bảng giá"
    ws.append(["Sản phẩm", "Đơn giá (VND)", "Đơn vị", "Hiệu lực từ", "Ghi chú"])
    ws.append(["Tủ lớn", 5000000, "Cái", "2026-01-01", "Tủ văn phòng loại A"])
    ws.append(["Bàn làm việc", 1200000, "Cái", "2026-01-01", "Bàn gỗ công nghiệp"])
    ws.append(["Ghế văn phòng", 850000, "Cái", "2026-01-01", "Ghế xoay có tựa lưng"])
    ws.append(["Máy tính xách tay", 15000000, "Cái", "2026-01-01", "Laptop 15 inch"])
    ws.append(["Màn hình máy tính", 4500000, "Cái", "2026-01-01", "Monitor 24 inch Full HD"])
    ws.append(["Giấy A4", 85000, "Ram", "2026-01-01", "500 tờ mỗi ram"])
    ws.append(["Mực máy in", 350000, "Hộp", "2026-01-01", "Tương thích HP LaserJet"])
    wb.save(os.path.join(HERE, "bang_gia.xlsx"))


if __name__ == "__main__":
    build()
