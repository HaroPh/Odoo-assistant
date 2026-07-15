import os, sys
sys.path.insert(0, os.path.join(os.path.dirname(__file__), "../../.."))
import pytest

FIX = os.path.join(os.path.dirname(__file__), "fixtures")


@pytest.fixture(scope="module")
def policy_docx():
    os.makedirs(FIX, exist_ok=True)
    path = os.path.join(FIX, "policy.docx")
    from docx import Document
    d = Document()
    d.add_heading("Chính sách hoàn hàng", level=1)
    d.add_heading("Mục 1 — Điều kiện", level=2)
    d.add_paragraph("Khách hàng có thể hoàn hàng trong vòng 30 ngày kể từ ngày mua. "
                    "Sản phẩm phải còn nguyên vẹn và có hóa đơn.")
    d.add_heading("Mục 2 — Ngoại lệ", level=2)
    d.add_paragraph("Hàng giảm giá không được hoàn trả. Hàng điện tử áp dụng bảo hành riêng.")
    d.save(path)
    return path


def test_parse_docx_yields_headings_and_body(policy_docx):
    from backend.src.rag.parse import parse_docx
    blocks = parse_docx(policy_docx)
    heads = [b["text"] for b in blocks if b["heading_level"]]
    assert "Chính sách hoàn hàng" in heads
    assert "Mục 2 — Ngoại lệ" in heads


def test_chunk_text_builds_section_path_and_does_not_straddle(policy_docx):
    from backend.src.rag.parse import parse_docx
    from backend.src.rag.chunking import chunk_text_blocks
    blocks = parse_docx(policy_docx)
    chunks = chunk_text_blocks(blocks, doc_id="policy", source_file=policy_docx)

    # every chunk has a breadcrumb rooted at the doc title
    assert all(c["section_path"].startswith("Chính sách hoàn hàng") for c in chunks)
    assert all(c["doc_title"] == "Chính sách hoàn hàng" for c in chunks)
    # the "30 ngày" fact and the "giảm giá" fact live in DIFFERENT sections (no straddle)
    s_30 = next(c["section_path"] for c in chunks if "30 ngày" in c["chunk_text"])
    s_sale = next(c["section_path"] for c in chunks if "giảm giá" in c["chunk_text"])
    assert "Mục 1" in s_30 and "Mục 2" in s_sale
    # chunk_index is sequential from 0
    assert [c["chunk_index"] for c in chunks] == list(range(len(chunks)))
    assert all(c["token_count"] > 0 for c in chunks)


def test_consecutive_same_level_headings_do_not_lose_content():
    # Bug 2026-07-15: 2 heading cùng cấp liên tiếp không body ở giữa → heading
    # trước bị pop mất, text heading sau chỉ còn trong section_path (không
    # được index). Guard mới: block heading thứ hai được giữ làm BODY.
    from backend.src.rag.chunking import chunk_text_blocks
    blocks = [
        {"text": "Điều 124. Trường hợp bị cưỡng chế", "heading_level": 2, "page": 64},
        {"text": "1. Người nộp thuế có tiền thuế nợ quá 90 ngày.", "heading_level": 2, "page": 64},
        {"text": "2. Người nộp thuế có tiền thuế nợ khi hết thời hạn gia hạn.", "heading_level": None, "page": 64},
    ]
    chunks = chunk_text_blocks(blocks, doc_id="d", source_file="f.pdf")
    bodies = " ".join(c["chunk_text"] for c in chunks)
    crumbs = " ".join(c["section_path"] for c in chunks)
    assert "quá 90 ngày" in bodies                    # khoản nằm trong chunk_text
    assert "Điều 124" in crumbs                       # heading thật giữ vai trò crumb


def test_two_headings_only_document_keeps_both_texts():
    # Trước fix: cả hai text biến mất (0 chunk). Sau fix: 1 chunk, text sau
    # là body, text trước là crumb — bất biến "không mất nội dung âm thầm".
    from backend.src.rag.chunking import chunk_text_blocks
    blocks = [
        {"text": "Chương I", "heading_level": 2, "page": 1},
        {"text": "NHỮNG QUY ĐỊNH CHUNG", "heading_level": 2, "page": 1},
    ]
    chunks = chunk_text_blocks(blocks, doc_id="d", source_file="f.pdf")
    assert len(chunks) == 1
    assert chunks[0]["chunk_text"] == "NHỮNG QUY ĐỊNH CHUNG"
    assert "Chương I" in chunks[0]["section_path"]


def test_index_text_composes_crumb_and_body():
    from backend.src.rag.chunking import index_text
    assert index_text("Điều 1. Phạm vi", "Nội dung khoản.") == "Điều 1. Phạm vi Nội dung khoản."
    assert index_text(None, "Chỉ có body.") == "Chỉ có body."
    assert index_text("", "Body.") == "Body."
