import os, sys
sys.path.insert(0, os.path.join(os.path.dirname(__file__), "../../.."))
import pytest


@pytest.fixture
def _mock_embed(monkeypatch):
    from backend.src.rag import ingest as ing
    monkeypatch.setattr(ing, "embed_texts",
                        lambda texts: [[0.01] * 1024 for _ in texts])


def _make_docx(path, body="Khách hàng có thể hoàn hàng trong 30 ngày."):
    from docx import Document
    d = Document()
    d.add_heading("Chính sách hoàn hàng", level=1)
    d.add_paragraph(body)
    d.save(path)


def test_ingest_inserts_then_skips_unchanged(clean_tables, _mock_embed, tmp_path):
    from backend.src.rag.ingest import ingest_path
    p = str(tmp_path / "policy.docx")
    _make_docx(p)

    s1 = ingest_path(p, conn=clean_tables)
    assert s1["ingested"] == 1 and s1["chunks"] >= 1
    n = clean_tables.execute("SELECT count(*) FROM rag_chunks").fetchone()[0]
    assert n == s1["chunks"]

    s2 = ingest_path(p, conn=clean_tables)         # unchanged → skip
    assert s2["ingested"] == 0 and s2["skipped"] == 1
    assert clean_tables.execute("SELECT count(*) FROM rag_chunks").fetchone()[0] == n


def test_ingest_replaces_changed_file_without_dupes(clean_tables, _mock_embed, tmp_path):
    from backend.src.rag.ingest import ingest_path
    p = str(tmp_path / "policy.docx")
    _make_docx(p, body="Hoàn hàng trong 30 ngày.")
    ingest_path(p, conn=clean_tables)
    _make_docx(p, body="Hoàn hàng trong 14 ngày. Áp dụng từ hôm nay với nhiều điều khoản hơn.")
    ingest_path(p, conn=clean_tables)              # changed → replace

    docs = clean_tables.execute("SELECT count(*) FROM rag_documents").fetchone()[0]
    assert docs == 1
    texts = clean_tables.execute("SELECT chunk_text FROM rag_chunks").fetchall()
    joined = " ".join(t[0] for t in texts)
    assert "14 ngày" in joined and "30 ngày" not in joined


def test_ingest_same_file_from_different_cwd_does_not_duplicate(
        clean_tables, _mock_embed, tmp_path, monkeypatch):
    # Bug: doc_id was os.path.relpath(path), which is relative to the
    # process's CWD at call time. Ingesting the SAME physical file via two
    # different cwd-relative spellings (e.g. running the ingest CLI from the
    # repo root vs. from backend/) must not create two rows for one document.
    from backend.src.rag.ingest import ingest_path
    sub = tmp_path / "seed"
    sub.mkdir()
    p = sub / "policy.docx"
    _make_docx(str(p))

    monkeypatch.chdir(tmp_path)
    s1 = ingest_path("seed/policy.docx", conn=clean_tables)
    assert s1["ingested"] == 1

    monkeypatch.chdir(sub)
    s2 = ingest_path("policy.docx", conn=clean_tables)      # same file, different cwd
    assert s2["ingested"] == 0 and s2["skipped"] == 1

    docs = clean_tables.execute("SELECT count(*) FROM rag_documents").fetchone()[0]
    assert docs == 1
    n = clean_tables.execute("SELECT count(*) FROM rag_chunks").fetchone()[0]
    assert n == s1["chunks"]


def test_ts_vector_is_populated(clean_tables, _mock_embed, tmp_path):
    from backend.src.rag.ingest import ingest_path
    p = str(tmp_path / "policy.docx")
    _make_docx(p)
    ingest_path(p, conn=clean_tables)
    nulls = clean_tables.execute("SELECT count(*) FROM rag_chunks WHERE ts_vector IS NULL").fetchone()[0]
    assert nulls == 0


def test_ingest_rolls_back_doc_on_embed_failure(clean_tables, monkeypatch, tmp_path):
    from backend.src.rag import ingest as ing
    from backend.src.rag.embed import EmbeddingError

    def _boom(texts):
        raise EmbeddingError("ollama down")

    monkeypatch.setattr(ing, "embed_texts", _boom)
    p = str(tmp_path / "policy.docx")
    from docx import Document
    d = Document()
    d.add_heading("Chính sách hoàn hàng", level=1)
    d.add_paragraph("Khách hàng có thể hoàn hàng trong 30 ngày.")
    d.save(p)

    with pytest.raises(Exception):
        ing.ingest_path(p, conn=clean_tables)
    # no orphan doc row, no chunks
    assert clean_tables.execute("SELECT count(*) FROM rag_documents").fetchone()[0] == 0
    assert clean_tables.execute("SELECT count(*) FROM rag_chunks").fetchone()[0] == 0


def test_ingest_indexes_section_path_but_keeps_chunk_text_body_only(
        clean_tables, monkeypatch, tmp_path):
    # Spec 2026-07-15 §3C: crumb vào chuỗi embed + ts_vector để tìm được theo
    # số Điều; cột chunk_text giữ body-only (hiển thị/citation nguyên trạng).
    from backend.src.rag import ingest as ing
    captured = []

    def _capture(texts):
        captured.extend(texts)
        return [[0.01] * 1024 for _ in texts]

    monkeypatch.setattr(ing, "embed_texts", _capture)

    from docx import Document
    p = str(tmp_path / "law.docx")
    d = Document()
    d.add_heading("Điều 99. Quy định đặc thù", level=1)
    d.add_paragraph("1. Nội dung khoản một về nghĩa vụ.")
    d.save(p)
    ing.ingest_path(p, conn=clean_tables)

    # Chuỗi embed chứa crumb
    assert any("Điều 99" in t for t in captured)
    # chunk_text trong DB vẫn body-only
    texts = [r[0] for r in clean_tables.execute(
        "SELECT chunk_text FROM rag_chunks").fetchall()]
    assert any("khoản một" in t for t in texts)
    assert all("Điều 99" not in t for t in texts)
    # ts_vector chứa token của crumb (cùng transform segment_vi hai phía)
    from backend.src.rag.ingest import segment_vi
    n = clean_tables.execute(
        "SELECT count(*) FROM rag_chunks "
        "WHERE ts_vector @@ plainto_tsquery('simple', %s)",
        (segment_vi("Điều 99"),)).fetchone()[0]
    assert n >= 1
